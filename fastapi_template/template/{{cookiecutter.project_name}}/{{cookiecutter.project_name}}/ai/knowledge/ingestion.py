"""Structure-preserving knowledge ingestion primitives.

Loaders deliberately return a small canonical document shape so providers can
be swapped without changing retrieval or business code.
"""

from __future__ import annotations

import hashlib
import ipaddress
import socket
from urllib.parse import urlparse
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


@dataclass(frozen=True, slots=True)
class KnowledgeDocument:
    """Normalized source document before chunking and indexing."""

    document_id: str
    source: str
    text: str
    metadata: dict[str, Any] = field(default_factory=dict)
    version: str = "1"


@dataclass(frozen=True, slots=True)
class KnowledgeChunk:
    """A retrieval unit with source and version provenance."""

    chunk_id: str
    document_id: str
    text: str
    metadata: dict[str, Any] = field(default_factory=dict)


def load_text(path: str | Path, *, metadata: dict[str, Any] | None = None) -> KnowledgeDocument:
    """Load UTF-8 text while preserving the source path."""
    source = str(Path(path))
    text = Path(path).read_text(encoding="utf-8")
    digest = hashlib.sha256(text.encode("utf-8")).hexdigest()
    return KnowledgeDocument(
        document_id=digest[:24],
        source=source,
        text=text,
        metadata=dict(metadata or {}),
        version=digest,
    )


def load_url(url: str, *, timeout_s: float = 10.0) -> KnowledgeDocument:
    """Load a URL through an explicit timeout and record its source."""
    import httpx
    import httpcore

    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"} or not parsed.hostname or parsed.username:
        raise ValueError("URL must use HTTP(S) and must not contain credentials")
    addresses = _resolve_public_addresses(parsed.hostname, parsed.port, parsed.scheme)
    if not addresses:
        raise ValueError("private, loopback, and non-global URL targets are blocked")
    pinned_address = str(sorted(addresses, key=str)[0])

    class PinnedBackend(httpcore.SyncBackend):
        """Connect to the address validated above while preserving TLS SNI."""

        def connect_tcp(
            self,
            host: str,
            port: int,
            timeout: float | None = None,
            local_address: str | None = None,
            socket_options: Any = None,
        ) -> Any:
            del host
            return super().connect_tcp(
                pinned_address,
                port,
                timeout,
                local_address,
                socket_options,
            )

    transport = httpx.HTTPTransport(trust_env=False, proxy=None)
    transport._pool = httpcore.ConnectionPool(  # type: ignore[attr-defined]
        network_backend=PinnedBackend(),
    )
    with httpx.Client(
        timeout=timeout_s,
        follow_redirects=False,
        transport=transport,
    ) as client:
        response = client.get(url)
    # Detect DNS changes across the request and fail closed for subsequent
    # retries. Proxies are disabled above so the request cannot bypass this
    # host validation through an environment-configured proxy.
    if _resolve_public_addresses(parsed.hostname, parsed.port, parsed.scheme) != addresses:
        raise ValueError("URL hostname changed during request")
    response.raise_for_status()
    if int(response.headers.get("content-length", "0") or 0) > 10_000_000:
        raise ValueError("URL response exceeds the 10 MB ingestion limit")
    if len(response.content) > 10_000_000:
        raise ValueError("URL response exceeds the 10 MB ingestion limit")
    text = response.text
    digest = hashlib.sha256(text.encode("utf-8")).hexdigest()
    return KnowledgeDocument(
        document_id=digest[:24],
        source=url,
        text=text,
        metadata={"content_type": response.headers.get("content-type", "")},
        version=digest,
    )


def _resolve_public_addresses(
    hostname: str,
    port: int | None,
    scheme: str,
) -> frozenset[ipaddress.IPv4Address | ipaddress.IPv6Address]:
    """Resolve every address and reject any non-public destination."""
    try:
        addresses = frozenset(
            ipaddress.ip_address(result[4][0])
            for result in socket.getaddrinfo(
                hostname,
                port or (443 if scheme == "https" else 80),
                type=socket.SOCK_STREAM,
            )
        )
    except (OSError, ValueError) as exc:
        raise ValueError("URL hostname could not be resolved") from exc
    if any(not address.is_global for address in addresses):
        raise ValueError("private, loopback, and non-global URL targets are blocked")
    return addresses


def load_pdf(path: str | Path, *, metadata: dict[str, Any] | None = None) -> KnowledgeDocument:
    """Extract text from a PDF while preserving page boundaries in metadata."""
    from pypdf import PdfReader

    source = str(Path(path))
    pages = [page.extract_text() or "" for page in PdfReader(source).pages]
    text = "\n\n".join(f"[Page {index}] {page}" for index, page in enumerate(pages, 1))
    digest = hashlib.sha256(text.encode("utf-8")).hexdigest()
    return KnowledgeDocument(
        document_id=digest[:24],
        source=source,
        text=text,
        metadata={**(metadata or {}), "page_count": len(pages), "format": "pdf"},
        version=digest,
    )


def load_docx(path: str | Path, *, metadata: dict[str, Any] | None = None) -> KnowledgeDocument:
    """Extract paragraph text from a DOCX without flattening source identity."""
    from docx import Document

    source = str(Path(path))
    text = "\n\n".join(paragraph.text for paragraph in Document(source).paragraphs if paragraph.text.strip())
    digest = hashlib.sha256(text.encode("utf-8")).hexdigest()
    return KnowledgeDocument(
        document_id=digest[:24],
        source=source,
        text=text,
        metadata={**(metadata or {}), "format": "docx"},
        version=digest,
    )


def chunk_document(
    document: KnowledgeDocument,
    *,
    max_chars: int = 1800,
    overlap: int = 180,
) -> list[KnowledgeChunk]:
    """Split on paragraph boundaries before falling back to character windows."""
    if max_chars <= 0 or overlap < 0 or overlap >= max_chars:
        raise ValueError("max_chars must be positive and overlap must be smaller")
    paragraphs = [part.strip() for part in document.text.split("\n\n") if part.strip()]
    chunks: list[KnowledgeChunk] = []
    current = ""
    for paragraph in paragraphs:
        if current and len(current) + len(paragraph) + 2 > max_chars:
            chunks.extend(_window_chunk(document, current, max_chars, overlap))
            current = current[-overlap:] if overlap else ""
        current = f"{current}\n\n{paragraph}".strip()
    if current:
        chunks.extend(_window_chunk(document, current, max_chars, overlap))
    return chunks


def _window_chunk(
    document: KnowledgeDocument,
    text: str,
    max_chars: int,
    overlap: int,
) -> list[KnowledgeChunk]:
    chunks: list[KnowledgeChunk] = []
    start = 0
    index = 0
    while start < len(text):
        end = min(len(text), start + max_chars)
        piece = text[start:end].strip()
        if piece:
            chunk_id = hashlib.sha256(
                f"{document.document_id}:{index}:{piece}".encode("utf-8")
            ).hexdigest()[:24]
            chunks.append(
                KnowledgeChunk(
                    chunk_id=chunk_id,
                    document_id=document.document_id,
                    text=piece,
                    metadata={
                        **document.metadata,
                        "source": document.source,
                        "version": document.version,
                        "chunk_index": index,
                    },
                )
            )
        if end >= len(text):
            break
        start = end - overlap
        index += 1
    return chunks


__all__ = [
    "KnowledgeChunk",
    "KnowledgeDocument",
    "chunk_document",
    "load_docx",
    "load_pdf",
    "load_text",
    "load_url",
]
